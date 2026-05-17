from __future__ import annotations

import io

from docx import Document
from hermes_api.services.docx import render_docx


def test_render_docx_returns_valid_bytes() -> None:
    minuta = (
        "[[CORPO]]\n"
        "PROCESSO Nº 0001234-56.2023.5.06.0020\n"
        "\n"
        "RECURSO DE REVISTA DA RECLAMADA\n"
        "\n"
        "TEMA - DANO MORAL\n"
        "\n"
        "Trata-se de recurso interposto pela reclamada.\n"
        "[[TRANSCRICAO1]]\n"
        "Transcrição do acórdão regional.\n"
        "[[CORPO]]\n"
        "Conheço do recurso e dou provimento.\n"
    )
    blob = render_docx(minuta)
    assert blob.startswith(b"PK")  # docx é zip
    doc = Document(io.BytesIO(blob))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("RECURSO DE REVISTA DA RECLAMADA" in t for t in texts)
    assert any("Conheço do recurso" in t for t in texts)


def test_repetitive_theme_audit_inserts_note() -> None:
    minuta = (
        "[[CORPO]]\n"
        "RECURSO DE REVISTA DA RECLAMADA\n"
        "\n"
        "HORAS EXTRAS\n"
        "\n"
        "Análise do tema.\n"
    )
    audit = {
        "capitulos": [
            {
                "titulo_minuta": "HORAS EXTRAS",
                "titulo_normalizado": "HORAS EXTRAS",
                "recurso": "RECURSO DE REVISTA DA RECLAMADA",
                "destacar": True,
                "nota_word": "Tema repetitivo identificado: Tema 99. Suspensão: ativa.",
            }
        ]
    }
    blob = render_docx(minuta, temas_repetitivos_audit=audit)
    doc = Document(io.BytesIO(blob))
    texts = [p.text for p in doc.paragraphs]
    assert any("Tema 99" in t for t in texts)


def test_transcricao_styles_have_progressive_indents() -> None:
    """Recuos esquerdos das Transcrições 1/2/3 devem progredir conforme docx.py."""
    import io
    from docx import Document
    from hermes_api.services.docx import render_docx

    blob = render_docx(
        "[[CORPO]]\nTEMA - X\n[[TRANSCRICAO1]]\na\n"
        "[[TRANSCRICAO2]]\nb\n[[TRANSCRICAO3]]\nc\n"
        "[[CORPO]]\nDISPOSITIVO\n"
    )
    d = Document(io.BytesIO(blob))
    styles = {s.name: s for s in d.styles}
    for name in ("Transcrição 1", "Transcrição 2", "Transcrição 3"):
        assert name in styles, f"estilo {name} ausente"

    def cm(name: str) -> float:
        return styles[name].paragraph_format.left_indent.cm

    assert abs(cm("Transcrição 1") - 4.5) < 0.01
    assert abs(cm("Transcrição 2") - 6.5) < 0.01
    assert abs(cm("Transcrição 3") - 7.5) < 0.01
    # first-line indent comum (1 cm)
    for name in ("Transcrição 1", "Transcrição 2", "Transcrição 3"):
        assert abs(styles[name].paragraph_format.first_line_indent.cm - 1) < 0.01
