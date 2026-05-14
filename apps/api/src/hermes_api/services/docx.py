"""Geração da minuta DOCX padrão TST.

Portado de scripts/generate_tst_docx.py. Recebe a minuta em Markdown/texto
estruturado (string) e devolve bytes do arquivo .docx. A validação de
anonimização final do script original foi removida deste porte (substituída
pelo pipeline de anonimização interno do Hermes em ``anonymizer.py``).
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_NAME = "Open Sans"
REPETITIVE_THEME_RED = RGBColor(255, 0, 0)
STYLE_MARKERS = {
    "[[CORPO]]": "Corpo",
    "[[ALERTA_VERMELHO]]": "Alerta Vermelho",
    "[[EMENTA]]": "Ementa",
    "[[TRANSCRICAO1]]": "Transcrição 1",
    "[[TRANSCRICAO2]]": "Transcrição 2",
    "[[TRANSCRICAO3]]": "Transcrição 3",
    "[[NOTA]]": "Notas de Rodapé",
}

LATIN_TERMS = [
    "quantum",
    "in casu",
    "data venia",
    "per se",
    "mutatis mutandis",
    "ex tunc",
    "ex nunc",
]

DECISION_PHRASES = [
    "conheço",
    "não conheço",
    "dou provimento",
    "dou-lhe provimento",
    "dou parcial provimento",
    "dou-lhe parcial provimento",
    "nego provimento",
    "nego-lhe provimento",
    "nego seguimento",
    "dou seguimento",
    "julgo prejudicado",
    "reputo prejudicado",
    "mantenho",
    "determino",
    "defiro",
    "indefiro",
]

FUNCTIONAL_PARTY_NAMES = {
    "reclamante": "Reclamante",
    "reclamada": "Reclamada",
    "reclamado": "Reclamado",
    "reclamantes": "Reclamantes",
    "reclamadas": "Reclamadas",
    "reclamados": "Reclamados",
    "embargante": "Embargante",
    "embargada": "Embargada",
    "embargado": "Embargado",
    "embargantes": "Embargantes",
    "embargadas": "Embargadas",
    "embargados": "Embargados",
    "agravante": "Agravante",
    "agravada": "Agravada",
    "agravado": "Agravado",
    "agravantes": "Agravantes",
    "agravadas": "Agravadas",
    "agravados": "Agravados",
    "recorrente": "Recorrente",
    "recorrentes": "Recorrentes",
    "recorrida": "Recorrida",
    "recorrido": "Recorrido",
    "recorridas": "Recorridas",
    "recorridos": "Recorridos",
}


def _set_font(run, size_pt: int, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _set_run_style(run, size_pt: int, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> None:
    _set_font(run, size_pt, bold)
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _set_style_font(style, size_pt: int, bold: bool = False) -> None:
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(size_pt)
    font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _set_doc_defaults(document: Document) -> None:
    styles = document.styles
    _set_style_font(styles["Normal"], 12)
    styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    section = document.sections[0]
    section.top_margin = Cm(4.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(1.0)

    body = styles.add_style("Corpo", 1)
    _set_style_font(body, 12)
    body.paragraph_format.first_line_indent = Cm(4.5)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.line_spacing = Pt(18)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    alert = styles.add_style("Alerta Vermelho", 1)
    _set_style_font(alert, 12, bold=True)
    alert.paragraph_format.first_line_indent = Cm(4.5)
    alert.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    alert.paragraph_format.line_spacing = Pt(18)
    alert.paragraph_format.space_before = Pt(0)
    alert.paragraph_format.space_after = Pt(0)
    alert.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading = styles.add_style("Corpo Título", 1)
    _set_style_font(heading, 12, bold=True)
    heading.paragraph_format.first_line_indent = Cm(4.5)
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading.paragraph_format.line_spacing = Pt(18)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ementa = styles.add_style("Ementa", 1)
    _set_style_font(ementa, 12)
    ementa.paragraph_format.left_indent = Cm(8)
    ementa.paragraph_format.first_line_indent = Cm(0)
    ementa.paragraph_format.line_spacing = 1
    ementa.paragraph_format.space_before = Pt(0)
    ementa.paragraph_format.space_after = Pt(0)
    ementa.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, left_cm in (("Transcrição 1", 4.5), ("Transcrição 2", 6.5), ("Transcrição 3", 7.5)):
        style = styles.add_style(name, 1)
        _set_style_font(style, 10)
        style.paragraph_format.left_indent = Cm(left_cm)
        style.paragraph_format.first_line_indent = Cm(1)
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note = styles.add_style("Notas de Rodapé", 1)
    _set_style_font(note, 10)
    note.paragraph_format.left_indent = Cm(0)
    note.paragraph_format.first_line_indent = Cm(0)
    note.paragraph_format.line_spacing = 1
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _strip_markdown(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^#{1,6}\s+", "", text)
    text = re.sub(r"^\*{1,3}(.*?)\*{1,3}$", r"\1", text)
    text = re.sub(r"^\*\*(.*?)\*\*$", r"\1", text)
    return text.strip()


def _is_heading(text: str) -> bool:
    stripped = _strip_markdown(text).strip()
    if not stripped:
        return False
    if stripped.startswith(("TEMA ", "I - ", "II - ", "III - ", "IV - ", "V - ")):
        return True
    if stripped in {
        "I - CONHECIMENTO",
        "II - MÉRITO",
        "REQUISITOS EXTRÍNSECOS DE ADMISSIBILIDADE",
        "CONHECIMENTO",
        "MÉRITO",
        "DISPOSITIVO",
    }:
        return True
    letters = [ch for ch in stripped if ch.isalpha()]
    return bool(letters) and stripped == stripped.upper() and len(stripped) <= 140


def _normalize_heading_key(text: str) -> str:
    return re.sub(r"\s+", " ", _strip_markdown(text)).strip().upper()


def _is_resource_heading(text: str) -> bool:
    stripped = _normalize_heading_key(text)
    stripped = re.sub(r"^[IVXLCDM]+\s*-\s*", "", stripped)
    return stripped.startswith(
        (
            "AGRAVO ",
            "AGRAVO DE ",
            "AGRAVO INTERNO ",
            "AGRAVO DE INSTRUMENTO ",
            "RECURSO ",
            "RECURSO DE ",
            "EMBARGOS ",
        )
    )


def _normalize_theme_heading(text: str) -> str:
    stripped = _strip_markdown(text).strip()
    if not stripped.upper().startswith("TEMA"):
        return text
    match = re.match(
        r"^(?P<prefix>TEMA)\s*(?:N[.ºO]?\s*)?(?P<num>\d+)?\s*[-–—]\s*(?P<desc>.+)$",
        stripped,
        flags=re.IGNORECASE,
    )
    if match:
        desc = match.group("desc").strip().replace(". ", " - ")
        return desc.upper()
    match = re.match(r"^(?P<prefix>TEMA)\s*[-–—]\s*(?P<desc>.+)$", stripped, flags=re.IGNORECASE)
    if match:
        desc = match.group("desc").strip().replace(". ", " - ")
        return desc.upper()
    return text


def _normalize_functional_party_names(text: str) -> str:
    role_pattern = "|".join(sorted((re.escape(k) for k in FUNCTIONAL_PARTY_NAMES), key=len, reverse=True))
    contextual = re.compile(
        rf"\b(?P<prefix>o|a|os|as|do|da|dos|das|ao|à|aos|às|pelo|pela|pelos|pelas|parte)\s+"
        rf"(?P<role>{role_pattern})\b",
        re.IGNORECASE,
    )

    def repl(match: re.Match[str]) -> str:
        role = FUNCTIONAL_PARTY_NAMES[match.group("role").lower()]
        return f"{match.group('prefix')} {role}"

    return contextual.sub(repl, text)


def _add_plain_runs_with_latin_italics(paragraph, text: str, *, size_pt: int, bold: bool, color: RGBColor | None = None) -> None:
    if not text:
        return
    latin_pattern = re.compile(
        r"\b(" + "|".join(re.escape(t) for t in sorted(LATIN_TERMS, key=len, reverse=True)) + r")\b",
        re.IGNORECASE,
    )
    decision_pattern = re.compile(
        r"\b(" + "|".join(re.escape(t) for t in sorted(DECISION_PHRASES, key=len, reverse=True)) + r")\b",
        re.IGNORECASE,
    )
    combined = re.compile("|".join([latin_pattern.pattern, decision_pattern.pattern]), re.IGNORECASE)
    pos = 0
    for m in combined.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            _set_run_style(run, size_pt, bold=bold, italic=False, color=color)
        token = m.group(0)
        token_lower = token.lower()
        is_latin = any(token_lower == t.lower() for t in LATIN_TERMS)
        is_decision = any(token_lower == t.lower() for t in DECISION_PHRASES)
        run = paragraph.add_run(token)
        _set_run_style(run, size_pt, bold=(bold or is_decision), italic=is_latin, color=color)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_style(run, size_pt, bold=bold, italic=False, color=color)


def _add_markdown_runs(
    paragraph,
    text: str,
    default_bold: bool = False,
    size_pt: int = 12,
    color: RGBColor | None = None,
) -> bool:
    text = _strip_markdown(text)
    if not text:
        return False
    pattern = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")
    used_manual_highlight = False
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            _add_plain_runs_with_latin_italics(
                paragraph, text[pos:match.start()],
                size_pt=size_pt, bold=default_bold, color=color,
            )
        token = match.group(0)
        if token.startswith("***") and token.endswith("***"):
            content = token[3:-3]
            run = paragraph.add_run(content)
            _set_run_style(run, size_pt, bold=True, italic=True, color=color)
            used_manual_highlight = True
        elif token.startswith("**") and token.endswith("**"):
            content = token[2:-2]
            run = paragraph.add_run(content)
            _set_run_style(run, size_pt, bold=True, italic=False, color=color)
            used_manual_highlight = True
        else:
            content = token[1:-1]
            run = paragraph.add_run(content)
            _set_run_style(run, size_pt, bold=False, italic=True, color=color)
            used_manual_highlight = True
        pos = match.end()
    if pos < len(text):
        _add_plain_runs_with_latin_italics(
            paragraph, text[pos:], size_pt=size_pt, bold=default_bold, color=color,
        )
    return used_manual_highlight


def _style_size(style_name: str) -> int:
    return 10 if style_name.startswith("Transcrição") or style_name == "Notas de Rodapé" else 12


def _add_paragraph(document: Document, text: str, current_style: str, color: RGBColor | None = None) -> bool:
    text = text.rstrip()
    text = text.replace("[...]", "(...)").replace("[…]", "(...)")
    heading = current_style == "Corpo" and _is_heading(text)
    if not current_style.startswith("Transcrição") and not heading:
        text = _normalize_functional_party_names(text)
    if not text:
        document.add_paragraph("", style=current_style)
        return False
    style_name = "Corpo Título" if heading else current_style
    paragraph = document.add_paragraph(style=style_name)
    if heading:
        text = _normalize_theme_heading(text)
    effective_color = color or (REPETITIVE_THEME_RED if current_style == "Alerta Vermelho" else None)
    return _add_markdown_runs(
        paragraph, text,
        default_bold=heading or current_style == "Alerta Vermelho",
        size_pt=_style_size(style_name),
        color=effective_color,
    )


def _add_blank_if_needed(document: Document, style: str = "Corpo") -> None:
    if not document.paragraphs:
        document.add_paragraph("", style=style)
        return
    if document.paragraphs[-1].text.strip():
        document.add_paragraph("", style=style)


def _add_heading_with_spacing(document: Document, text: str, color: RGBColor | None = None) -> bool:
    _add_blank_if_needed(document, "Corpo")
    used = _add_paragraph(document, text, "Corpo", color=color)
    _add_blank_if_needed(document, "Corpo")
    return used


def _add_core_properties(document: Document) -> None:
    core = document.core_properties
    core.author = "Hermes TST"
    core.subject = "Minuta judicial formatada conforme padrão TST"
    core.keywords = "TST, Open Sans, minuta, voto, decisão"


def _enforce_update_fields(document: Document) -> None:
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _load_repetitive_annotations(audit: dict[str, Any] | None) -> dict[str, dict]:
    if not audit:
        return {"by_composite": {}, "by_title": {}}
    by_composite: dict[str, dict] = {}
    candidates_by_title: dict[str, list[dict]] = {}
    for item in audit.get("capitulos", []):
        if not item.get("destacar"):
            continue
        title = item.get("titulo_normalizado") or item.get("titulo_minuta")
        note = item.get("nota_word")
        if title and note:
            title_key = _normalize_heading_key(title)
            resource_key = _normalize_heading_key(item.get("recurso") or "")
            by_composite[f"{resource_key}\0{title_key}"] = item
            candidates_by_title.setdefault(title_key, []).append(item)
    by_title = {tk: items[0] for tk, items in candidates_by_title.items() if len(items) == 1}
    return {"by_composite": by_composite, "by_title": by_title}


def _find_repetitive_annotation(annotations: dict[str, dict], resource_title: str | None, heading: str) -> dict | None:
    title_key = _normalize_heading_key(heading)
    resource_key = _normalize_heading_key(resource_title or "")
    return annotations["by_composite"].get(f"{resource_key}\0{title_key}") or annotations["by_title"].get(title_key)


def render_docx(
    minuta_text: str,
    *,
    temas_repetitivos_audit: dict[str, Any] | None = None,
) -> bytes:
    """Renderiza a minuta em bytes do DOCX padrão TST."""
    document = Document()
    _set_doc_defaults(document)
    _add_core_properties(document)
    _enforce_update_fields(document)
    annotations = _load_repetitive_annotations(temas_repetitivos_audit)

    current_style = "Corpo"
    current_resource_title: str | None = None
    transcription_had_highlights = False

    for raw_line in minuta_text.splitlines():
        line = raw_line.rstrip()
        marker = line.strip().upper()
        if marker in STYLE_MARKERS:
            next_style = STYLE_MARKERS[marker]
            if next_style.startswith("Transcrição") and current_style != next_style:
                _add_paragraph(document, "", "Corpo")
                transcription_had_highlights = False
            if next_style == "Corpo" and current_style.startswith("Transcrição"):
                _add_paragraph(document, "", "Corpo")
                if transcription_had_highlights:
                    _add_paragraph(document, "(destaques acrescidos)", "Corpo")
                    _add_paragraph(document, "", "Corpo")
            current_style = next_style
            continue
        if not line.strip():
            continue
        if current_style == "Corpo" and _is_heading(line):
            rendered_heading = _normalize_theme_heading(line)
            annotation = None if _is_resource_heading(line) else _find_repetitive_annotation(
                annotations, current_resource_title, rendered_heading,
            )
            used_manual = _add_heading_with_spacing(
                document, line, color=REPETITIVE_THEME_RED if annotation else None,
            )
            if _is_resource_heading(line):
                current_resource_title = _normalize_heading_key(rendered_heading)
            if annotation:
                _add_paragraph(document, annotation["nota_word"], "Corpo")
                _add_blank_if_needed(document, "Corpo")
        else:
            used_manual = _add_paragraph(document, line, current_style)
        if current_style.startswith("Transcrição") and used_manual:
            transcription_had_highlights = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
