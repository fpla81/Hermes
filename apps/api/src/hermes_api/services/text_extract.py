"""Extração de texto de peças preparadas (txt/md/html/pdf/docx).

Portado de scripts/build_analysis_packets.py + scripts/check_resource_texts.py.
PyMuPDF e python-docx são opcionais: extração desses tipos só roda se as libs
estiverem instaladas; caso contrário devolve string vazia + erro descritivo.
"""

from __future__ import annotations

import html
import re
from dataclasses import dataclass
from html.parser import HTMLParser

SUPPORTED_SUFFIXES = (".txt", ".md", ".html", ".htm", ".pdf", ".docx")
_BLOCK_TAGS = {
    "address", "article", "aside", "blockquote", "br", "dd", "div", "dl",
    "dt", "figcaption", "figure", "footer", "h1", "h2", "h3", "h4", "h5",
    "h6", "header", "hr", "li", "main", "nav", "ol", "p", "pre",
    "section", "table", "tbody", "td", "tfoot", "th", "thead", "tr", "ul",
}


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=False)
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript"}:
            self.skip_depth += 1
            return
        if tag in _BLOCK_TAGS:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript"} and self.skip_depth:
            self.skip_depth -= 1
            return
        if tag in _BLOCK_TAGS:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    def handle_entityref(self, name: str) -> None:
        if not self.skip_depth:
            self.parts.append(f"&{name};")

    def handle_charref(self, name: str) -> None:
        if not self.skip_depth:
            self.parts.append(f"&#{name};")

    def text(self) -> str:
        return "".join(self.parts)


def normalize_spaces(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def strip_html(raw: str) -> str:
    if not re.search(r"<[a-zA-Z][^>]*>", raw):
        return raw
    parser = _TextExtractor()
    parser.feed(raw)
    return html.unescape(parser.text())


@dataclass
class ExtractedText:
    text: str
    error: str | None = None


def extract_text(content: bytes, filename: str) -> ExtractedText:
    """Devolve texto plano normalizado a partir do conteúdo bruto da peça.

    A escolha do extrator usa a extensão de ``filename``.
    """
    suffix = _suffix_of(filename)
    if suffix in {".txt", ".md", ".html", ".htm"}:
        decoded = content.decode("utf-8", errors="replace")
        return ExtractedText(text=normalize_spaces(strip_html(decoded)))
    if suffix == ".pdf":
        try:
            import fitz  # type: ignore[import-untyped]
        except ImportError:
            return ExtractedText(text="", error="PyMuPDF indisponível para extrair texto do PDF")
        try:
            with fitz.open(stream=content, filetype="pdf") as doc:
                pages = "\n".join(page.get_text("text") for page in doc)
            return ExtractedText(text=normalize_spaces(pages))
        except Exception as exc:  # noqa: BLE001
            return ExtractedText(text="", error=f"falha ao ler PDF: {exc}")
    if suffix == ".docx":
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError:
            return ExtractedText(text="", error="python-docx indisponível para extrair texto do DOCX")
        try:
            import io

            doc = Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return ExtractedText(text=normalize_spaces("\n".join(parts)))
        except Exception as exc:  # noqa: BLE001
            return ExtractedText(text="", error=f"falha ao ler DOCX: {exc}")
    return ExtractedText(text="", error=f"tipo de arquivo não suportado: {suffix}")


def _suffix_of(filename: str) -> str:
    dot = filename.rfind(".")
    return filename[dot:].lower() if dot >= 0 else ""
